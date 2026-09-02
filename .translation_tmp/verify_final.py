# -*- coding: utf-8 -*-
import zipfile, re, io, html, shutil, os
import xml.etree.ElementTree as ET

src = r"d:\Learning_test\backup3\ZhiTuPoJu\.translation_tmp\paper_translated.docx"
z = zipfile.ZipFile(src)
doc = z.read("word/document.xml").decode("utf-8")
ET.fromstring(doc)
print("XML valid")

ACC = re.compile(r"[àâçéèêëîïôùûüÿœæÀÂÇÉÈÊËÎÏÔÙÛÜŸŒÆ]")
fr = []
for m in re.finditer(r"<(?:ns0|ns9):t(?:\s[^>]*)?>(.*?)</(?:ns0|ns9):t>", doc, re.S):
    t = html.unescape(m.group(1))
    if ACC.search(t) and not re.fullmatch(r"[\d\s.,;:()%=\+\-−÷×≤≥<>/|_^]*", t):
        fr.append(t.strip()[:60])
print("accented non-formula nodes:", len(fr))
for f in fr:
    print(" -", repr(f))

cjk = []
for m in re.finditer(r"<ns0:t(?: xml:space=\"preserve\")?>(.*?)</ns0:t>", doc, re.S):
    t = html.unescape(m.group(1))
    if re.search(r"[\u4e00-\u9fff]", t):
        cjk.append(t.strip()[:60])
print("CJK nodes:", len(cjk))

import docx
d = docx.Document(src)
paras = [p.text for p in d.paragraphs if p.text.strip()]
print("paragraphs:", len(paras), "| tables:", len(d.tables))
print("Related Works:", any("Related Works" in p for p in paras))
print("Modification:", any("Modification:" in p for p in paras))

dst = r"d:\Learning_test\backup3\ZhiTuPoJu\paper_EN.docx"
if os.path.exists(dst):
    os.remove(dst)
shutil.copy(src, dst)
print("final output:", dst, os.path.getsize(dst))
