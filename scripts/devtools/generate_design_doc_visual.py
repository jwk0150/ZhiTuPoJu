# -*- coding: utf-8 -*-
"""
生成《作品设计实现方案》视觉精简版（参照善泽生物创业计划书体例）：
- 每页页眉角标装饰 + 页脚深蓝项目条
- 白底、青绿强调、大量留白
- 内容精炼，约 25–30 页量级，而非十万字论文体

运行：python scripts/devtools/generate_design_doc_visual.py
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFont, ImageFilter

ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent / "design_doc"
ASSETS = HERE / "assets"
OUT = ROOT / "作品设计实现方案_提交版.docx"
OUT_ALT = ROOT / "作品设计实现方案_视觉版.docx"

TEAL = RGBColor(0x2A, 0x8F, 0x9A)
TEAL_DARK = RGBColor(0x1A, 0x5C, 0x66)
NAVY = "1A3A4A"
GOLD = RGBColor(0xC9, 0xA8, 0x6A)
GRAY = RGBColor(0x3A, 0x4A, 0x50)
FOOTER_NAME = "执图破局 作品设计实现方案 · XH-202621"


def set_run(run, *, name="微软雅黑", size=11, bold=False, color=None, en=None):
    en = en or name
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), name)


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in kwargs.items():
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def configure_page(section, *, cover=False):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    if cover:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
    else:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)


def add_page_chrome(section):
    """每页：右上角装饰 + 页脚深蓝项目条。"""
    corner = ASSETS / "page-corner.png"
    header = section.header
    header.is_linked_to_previous = False
    # clear existing
    for p in list(header.paragraphs):
        p.clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    if corner.exists():
        run = hp.add_run()
        run.add_picture(str(corner), width=Cm(3.2))

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    table = footer.add_table(rows=1, cols=2, width=Cm(16.8))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    c0, c1 = table.rows[0].cells
    shade_cell(c0, NAVY)
    shade_cell(c1, NAVY)
    set_cell_margins(c0, top=70, bottom=70, left=140, right=40)
    set_cell_margins(c1, top=70, bottom=70, left=40, right=140)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run(FOOTER_NAME)
    set_run(r0, size=9, color=RGBColor(255, 255, 255), name="微软雅黑")
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p1.add_run()
    set_run(run, size=9, color=RGBColor(255, 255, 255), name="微软雅黑")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def h1(doc, cn: str, en: str = ""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(cn)
    set_run(r, name="微软雅黑", size=18, bold=True, color=TEAL)
    if en:
        r2 = p.add_run("  " + en)
        set_run(r2, name="Arial", size=12, bold=True, color=TEAL_DARK, en="Arial")
    return p


def h2(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, name="微软雅黑", size=13, bold=True, color=TEAL)
    return p


def body(doc, text: str, *, first_indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.35
    pf.space_after = Pt(6)
    if first_indent:
        pf.first_line_indent = Cm(0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # simple highlight: wrap **key** in teal bold
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if i % 2 == 1:
            set_run(run, size=11, bold=True, color=TEAL, name="微软雅黑")
        else:
            set_run(run, size=11, color=GRAY, name="微软雅黑")
    return p


def bullets(doc, items: list[str]):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run("●  " + it)
        set_run(r, size=11, color=GRAY, name="微软雅黑")


def tip_card(doc, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F3FAFB")
    set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    set_run(r, size=11, bold=True, color=TEAL, name="微软雅黑")
    r2 = p.add_run(text)
    set_run(r2, size=10.5, color=GRAY, name="微软雅黑")
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def font_path(cands, size):
    for p in cands:
        if Path(p).exists():
            try:
                return ImageFont.truetype(p, size=size)
            except Exception:
                pass
    return ImageFont.load_default()


def compose_cover():
    bg_src = Path(r"C:\Users\Ibiza\.cursor\projects\c-Users-Ibiza-Desktop-project\assets\zhitu-cover-shanze-style.png")
    if not bg_src.exists():
        bg_src = ASSETS / "cover-bg.png"
    W, H = 1240, 1754
    bg = Image.open(bg_src).convert("RGBA").resize((W, H), Image.Resampling.LANCZOS)
    canvas = Image.new("RGBA", (W, H), (255, 255, 255, 255))
    canvas.alpha_composite(bg)
    draw = ImageDraw.Draw(canvas)

    yahei = [r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\msyhbd.ttc"]
    heiti = [r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\msyhbd.ttc"]
    f_brand_cn = font_path(heiti, 28)
    f_brand_en = font_path(yahei, 14)
    f_title = font_path(heiti, 48)
    f_en = font_path(yahei, 18)
    f_sub = font_path(heiti, 26)
    f_meta = font_path(yahei, 22)
    f_dash = font_path(heiti, 24)

    # top-right brand
    logo = Image.open(ASSETS / "logo.png").convert("RGBA").resize((72, 72), Image.Resampling.LANCZOS)
    canvas.alpha_composite(logo, (W - 220, 56))
    draw.text((W - 140, 68), "执图破局", font=f_brand_cn, fill=(42, 92, 98, 255))
    draw.text((W - 140, 104), "ZHITU POJU", font=f_brand_en, fill=(90, 130, 138, 255))

    # bottom title block
    def center(text, y, fnt, fill):
        b = draw.textbbox((0, 0), text, font=fnt)
        tw = b[2] - b[0]
        draw.text(((W - tw) / 2, y), text, font=fnt, fill=fill)

    center("执图破局", 1180, f_title, (26, 72, 80, 255))
    center("Zhitu Pojue · Digital Talent Intelligence", 1255, f_en, (80, 120, 128, 255))
    center("多源异构数据驱动的岗位—能力知识图谱", 1310, f_sub, (40, 70, 78, 255))
    center("动态构建与智能匹配系统", 1355, f_sub, (40, 70, 78, 255))
    center("—— 作品设计实现方案", 1420, f_dash, (42, 143, 154, 255))
    center("赛题 XH-202621  ·  科大讯飞  ·  2026 揭榜挂帅", 1485, f_meta, (90, 110, 118, 255))
    center("河南工业大学 · 人工智能与大数据学院", 1530, f_meta, (90, 110, 118, 255))

    out = ASSETS / "cover-visual.png"
    canvas.convert("RGB").save(out, "PNG")
    return out


def build_toc(doc):
    h1(doc, "目录", "TABLE OF CONTENTS")
    entries = [
        ("一、执行总结", "EXECUTIVE SUMMARY"),
        ("二、项目背景", "PROJECT BACKGROUND"),
        ("三、核心技术", "CORE TECHNOLOGY"),
        ("四、系统架构", "SYSTEM ARCHITECTURE"),
        ("五、创新点", "INNOVATION"),
        ("六、测试与演示", "TESTING & DEMO"),
        ("七、总结与展望", "CONCLUSION"),
    ]
    for cn, en in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(f"{cn}    {en}")
        set_run(r, size=13, bold=True, color=TEAL, name="微软雅黑")
    page_break(doc)


def build_content(doc):
    # 1 执行总结
    h1(doc, "一、执行总结", "EXECUTIVE SUMMARY")
    h2(doc, "1.1 目的")
    body(
        doc,
        "本方案面向赛题 **XH-202621**，交付可运行的「执图破局」系统与可评测证据链。"
        "目标是让岗位—能力关系可洞察、新兴岗位可发现、人岗匹配可解释。",
    )
    h2(doc, "1.2 背景")
    body(
        doc,
        "传统招聘以关键词检索为主，岗位别名与能力组合变化快，导致发现滞后、匹配黑盒。"
        "平台以多源 JD 语料与知识图谱为底座，形成发现—演化—匹配闭环。",
    )
    h2(doc, "1.3 产品与服务")
    bullets(
        doc,
        [
            "数字人才地图：岗位/能力关系可视化",
            "新岗位发现：真实发现 + 未来预测双线",
            "能力演化：来源路径与差分复盘",
            "人岗匹配：五维评分与缺口行动清单",
            "智能问答：RAG 检索增强与幻觉防控",
        ],
    )
    tip_card(doc, "一句话定位", "用图谱与数据，把求职、研究与决策连成清晰路径。")
    page_break(doc)

    h2(doc, "1.4 项目优势")
    bullets(
        doc,
        [
            "**可回放**：发现与匹配结果可追溯到语料与规则版本",
            "**可解释**：五维匹配给出缺口与下一步动作",
            "**可演示**：一屏座舱前端，评委可快速走通主路径",
            "**可评测**：门控、差分、幻觉防控均有验收口径",
        ],
    )
    h2(doc, "1.5 技术路线")
    body(
        doc,
        "数据层使用 **PostgreSQL** 多库视图；服务层以 **FastAPI** 编排；"
        "推理增强接入 **DeepSeek**；前端以一屏模块化座舱呈现发现/演化/匹配结果。",
    )
    h2(doc, "1.6 团队与单位")
    body(
        doc,
        "申报单位：河南工业大学 · 人工智能与大数据学院。"
        "核心成员：傅英淮、李帅 等。发榜单位：科大讯飞。",
        first_indent=False,
    )
    page_break(doc)

    # 2 背景
    h1(doc, "二、项目背景", "PROJECT BACKGROUND")
    h2(doc, "2.1 问题定义")
    bullets(
        doc,
        [
            "岗位标题碎片化，同一能力组合对应多个别名",
            "新兴岗位信号淹没在海量 JD 中，难以及时识别",
            "匹配结果缺少能力缺口与行动建议，难以落地",
        ],
    )
    h2(doc, "2.2 解决思路")
    body(
        doc,
        "以「岗位—能力」为最小关系单元，先做可信入库，再做发现与演化，最后落到可解释匹配。"
        "预测岗位与真实发现分流，避免把前瞻信号误当作已招承诺。",
    )
    tip_card(
        doc,
        "赛题对齐",
        "直接服务揭榜挂帅赛题对知识图谱构建、动态更新与智能匹配的要求，强调可复现与可演示。",
    )
    page_break(doc)

    # 3 核心技术
    h1(doc, "三、核心技术", "CORE TECHNOLOGY")
    h2(doc, "3.1 技术总览")
    body(doc, "五项可定位创新（I1—I5）构成平台能力主轴：")
    bullets(
        doc,
        [
            "I1 入库门控：质量/合规/去重",
            "I2 新兴岗位发现：真实发现与预测分流",
            "I3 能力演化差分：可复算路径",
            "I4 五维人岗匹配：缺口可解释",
            "I5 RAG 幻觉防控：检索约束与拒答策略",
        ],
    )
    page_break(doc)

    h2(doc, "3.2 I1 入库门控")
    body(
        doc,
        "多源 JD 进入图谱前，经过字段完备性、噪声清洗与去重校验。"
        "门控失败样本可回看，保证后续发现与匹配建立在可信语料上。",
    )
    h2(doc, "3.3 I2 新兴岗位发现")
    body(
        doc,
        "真实发现线：聚类已在招聘文本中稳定出现的新岗位组合。"
        "预测线：交汇能力推演尚未固化的标题，必须附不确定性说明。",
    )
    tip_card(doc, "读法提醒", "预测岗是前瞻信号，不是录用保证；决策前必读不确定性模块。")
    page_break(doc)

    h2(doc, "3.4 I3 能力演化差分")
    body(
        doc,
        "以相邻已有岗为源，计算能力增减与迁移成本，输出主路径/次路径与必补清单，"
        "支持「从哪来、还差什么」的复盘。",
    )
    h2(doc, "3.5 I4 五维人岗匹配")
    body(
        doc,
        "从技能重合、职责覆盖、行业迁移、经验层次、成长潜力五维评分，"
        "并生成优先缺口与本周可做行动。",
    )
    page_break(doc)

    h2(doc, "3.6 I5 RAG 幻觉防控")
    body(
        doc,
        "问答与解读优先引用已入库证据；证据不足时降级或拒答，"
        "避免把模型推测写成事实结论。",
    )
    h2(doc, "3.7 关键能力")
    bullets(
        doc,
        [
            "一屏模块化详情：左导航 / 中主图 / 右解读同高对齐",
            "供需、趋势、雷达等可视化与行动清单联动",
            "简历对照报告：把窗口翻译成可补能力项",
        ],
    )
    page_break(doc)

    # 4 架构
    h1(doc, "四、系统架构", "SYSTEM ARCHITECTURE")
    h2(doc, "4.1 总体分层")
    bullets(
        doc,
        [
            "数据层：PostgreSQL（岗位/能力/匹配相关库与视图）",
            "服务层：FastAPI 路由 + 异步任务/批处理",
            "智能层：规则引擎 + DeepSeek 增强",
            "表现层：静态前端座舱（发现/地图/匹配/仓库）",
        ],
    )
    h2(doc, "4.2 关键链路")
    body(
        doc,
        "列表 → 详情模块切换 → 洞察栏行动 → 简历对照 / 收藏回看。"
        "每条链路均可在演示中 3 分钟内走通。",
    )
    tip_card(doc, "部署口径", "本地开发：后端 8000，前端静态服务；演示环境可同构迁移。")
    page_break(doc)

    h2(doc, "4.3 数据与安全")
    bullets(
        doc,
        [
            "语料入库留痕，支持版本对照",
            "用户简历与收藏数据隔离存储",
            "对外展示区分真实发现与预测，降低误读风险",
        ],
    )
    page_break(doc)

    # 5 创新
    h1(doc, "五、创新点", "INNOVATION")
    h2(doc, "5.1 创新矩阵")
    body(doc, "创新不追求概念堆叠，而强调**可定位、可验收、可演示**：")
    bullets(
        doc,
        [
            "发现可回放：样本与规则版本可追溯",
            "演化可复算：差分结果可重跑核对",
            "匹配可解释：五维拆解 + 行动清单",
            "预测可降噪：不确定性模块强制露出",
            "问答可约束：RAG 证据优先与拒答",
        ],
    )
    page_break(doc)

    h2(doc, "5.2 与同类方案差异")
    body(
        doc,
        "相较纯关键词招聘搜索或静态职业百科，本系统把「新岗位出现—能力迁移—人岗缺口」"
        "串成同一套图谱语言，评委可从同一岗位详情页同时看到证据、路径与行动。",
    )
    page_break(doc)

    # 6 测试
    h1(doc, "六、测试与演示", "TESTING & DEMO")
    h2(doc, "6.1 验收要点")
    bullets(
        doc,
        [
            "门控：脏数据拦截与日志可查",
            "发现：真实/预测分流展示正确",
            "匹配：五维分数与缺口一致",
            "前端：三列同高、模块切换无大面积留白",
            "问答：无证据时不硬编结论",
        ],
    )
    h2(doc, "6.2 演示脚本（建议 5 分钟）")
    bullets(
        doc,
        [
            "30s：封面定位与赛题编号",
            "90s：新岗位发现（真实岗详情 + 供需）",
            "90s：能力演化路径",
            "90s：人岗匹配与缺口行动",
            "30s：创新点与评测证据收束",
        ],
    )
    page_break(doc)

    # 7 总结
    h1(doc, "七、总结与展望", "CONCLUSION")
    body(
        doc,
        "执图破局以可信语料为基、以图谱关系为核、以可解释匹配为用，"
        "形成面向揭榜挂帅赛题的完整作品闭环。后续将持续扩充语料覆盖、"
        "强化演化评测集，并优化演示动线与报告导出体验。",
    )
    tip_card(
        doc,
        "提交清单",
        "可运行系统 + 本设计实现方案 + 演示脚本/录屏 + 关键接口与评测记录。",
    )


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    # ensure corner exists
    corner_src = Path(r"C:\Users\Ibiza\.cursor\projects\c-Users-Ibiza-Desktop-project\assets\zhitu-page-corner.png")
    if corner_src.exists():
        (ASSETS / "page-corner.png").write_bytes(corner_src.read_bytes())

    print("合成封面…")
    cover = compose_cover()

    doc = Document()
    # cover section
    configure_page(doc.sections[0], cover=True)
    for part in (doc.sections[0].header, doc.sections[0].footer):
        part.is_linked_to_previous = False
        if part.paragraphs:
            part.paragraphs[0].clear()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(cover), width=Cm(21.0), height=Cm(29.7))

    # body section
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_sec, cover=False)
    add_page_chrome(body_sec)

    build_toc(doc)
    build_content(doc)

    saved = OUT_ALT
    try:
        doc.save(OUT)
        saved = OUT
    except PermissionError:
        print("提交版文件被占用（请先关闭 Word），改存视觉版…")
        doc.save(OUT_ALT)
        saved = OUT_ALT
    except OSError as e:
        fallback = ROOT / "作品设计实现方案_视觉精简版.docx"
        print("保存失败，改用：", fallback, e)
        doc.save(fallback)
        saved = fallback
    print("已生成：", saved)
    print("体积：", round(saved.stat().st_size / 1024, 1), "KB")
    print("提示：Word 中更新域可刷新页脚页码。")


if __name__ == "__main__":
    main()
