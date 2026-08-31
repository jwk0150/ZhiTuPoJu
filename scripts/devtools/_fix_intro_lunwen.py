# -*- coding: utf-8 -*-
"""Rewrite intro: problem first, then introduce platform. Keep font style."""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

SRC = Path(r"C:\Users\Ibiza\Desktop\lunwen.docx")
DST = Path(r"C:\Users\Ibiza\Desktop\lunwen.docx")

NEW_TEXT = (
    "当前就业市场面临技术栈迭代明显快于人才培养周期的结构性矛盾："
    "企业侧招不到匹配人才，劳动者侧又难以看清能力成长路径。"
    "公开招聘文本本是观察岗位能力需求的重要信号，却普遍存在发布时滞、模板化复制与表述噪声，"
    "直接污染下游知识抽取与匹配结果；传统招聘又多依赖关键词比对，"
    "难以刻画岗位能力结构的动态演化。与此同时，大模型在岗位定义生成与诊断对话中易出现证据不足的幻觉，"
    "削弱结论可信度。针对上述问题，本文设计并实现面向就业市场的岗位—能力知识图谱动态构建与智能匹配系统「执图破局」。"
    "系统以岗位与能力为核心实体，融合多源数据采集、Neo4j图谱构建、RAG推理与幻觉防控、多维度人岗匹配算法，"
    "实现新兴岗位发现、能力动态更新、技能点级全景可视化及简历解析诊断闭环。"
    "在线评测中，JD解析、简历提取与匹配准确率均达90%以上，可支撑企业招聘、职业规划与人才态势研判。"
)


def copy_rpr(src_run, dst_run):
    """Copy character formatting from src_run to dst_run if present."""
    src_rPr = src_run._element.find(qn("w:rPr"))
    if src_rPr is None:
        # fallback: bold if source was bold
        if src_run.bold:
            dst_run.bold = True
        return
    dst_rPr = dst_run._element.get_or_add_rPr()
    # clear existing children
    for child in list(dst_rPr):
        dst_rPr.remove(child)
    for child in src_rPr:
        dst_rPr.append(deepcopy(child))


def main():
    doc = Document(str(SRC))
    target = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("平台概述") or ("执图破局" in t and "岗位—能力知识图谱" in t and "平台概述" in t):
            target = p
            break
        if t.startswith("平台概述:") or t.startswith("平台概述："):
            target = p
            break
    if target is None:
        # fallback: first paragraph containing 执图破局 after 引言
        seen_intro = False
        for p in doc.paragraphs:
            if p.text.strip().startswith("一、引言") or "引言" in p.text.strip()[:6]:
                seen_intro = True
                continue
            if seen_intro and "执图破局" in p.text:
                target = p
                break
    if target is None:
        raise SystemExit("target paragraph not found")

    # remember first run formatting
    template_run = target.runs[0] if target.runs else None

    # clear all runs
    for r in list(target.runs):
        r._element.getparent().remove(r._element)

    # also clear leftover text nodes directly under p (rare)
    new_run = target.add_run(NEW_TEXT)
    if template_run is not None:
        copy_rPr(template_run, new_run)
    else:
        new_run.bold = True

    # ensure east-asia font stays 宋体-ish if missing
    rPr = new_run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "宋体")
        rFonts.set(qn("w:cs"), "Times New Roman")

    doc.save(str(DST))
    print("updated:", DST)
    print("preview:", NEW_TEXT[:80], "...")


if __name__ == "__main__":
    main()
