# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import zipfile
from lxml import etree

path = Path(r"C:\Users\Ibiza\Desktop\lunwen - 副本.docx")
out = Path(r"C:\Users\Ibiza\Desktop\project\挑战杯\scripts\devtools\_lunwen_dump.txt")
doc = Document(str(path))

lines = []
lines.append(f"PARAS={len(doc.paragraphs)} TABLES={len(doc.tables)} SECTIONS={len(doc.sections)}")
lines.append("--- STYLES used ---")
styles = sorted({(p.style.name if p.style else "?") for p in doc.paragraphs})
lines.append(", ".join(styles))
lines.append("--- BODY ---")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    text = p.text.replace("\r", " ").replace("\n", " ").strip()
    if not text:
        # keep heading empties? skip blank
        continue
    mark = f"[{style}] " if ("Heading" in style or style.startswith("标题") or "Title" in style) else ""
    # first run bold?
    bold_prefix = ""
    if p.runs:
        r0 = p.runs[0]
        if r0.bold and r0.text.strip():
            bold_prefix = "{B}"
    lines.append(f"{i:04d}{bold_prefix} {mark}{text}")

# comments from word/comments.xml
lines.append("\n--- COMMENTS ---")
with zipfile.ZipFile(path) as z:
    names = z.namelist()
    lines.append("zip entries with comment: " + ", ".join(n for n in names if "comment" in n.lower()))
    if "word/comments.xml" in names:
        root = etree.fromstring(z.read("word/comments.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for c in root.findall("w:comment", ns):
            cid = c.get(qn("w:id"))
            author = c.get(qn("w:author"))
            texts = [t.text or "" for t in c.findall(".//w:t", ns)]
            lines.append(f"[{cid}] {author}: {''.join(texts)}")

out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out, "lines", len(lines))
