# -*- coding: utf-8 -*-
"""用户中心 — 业务逻辑层
负责：简历解析编排、AI 职业访谈、岗位匹配计算、画像生成
"""
from __future__ import annotations

import json
import math
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from backend.db import SessionLocal
from backend.llm import deepseek
from backend.models.user_profile import CareerReport, Resume, UserProfile, UserSkill

# ============================================================
# 文本提取（直接复用 matching/service.py 的成熟实现）
# ============================================================


def extract_resume_text(filepath: str, file_type: str) -> str:
    """根据文件类型提取文本，复用 matching 服务的实现"""
    ext = file_type.lower().lstrip(".")
    filepath_obj = Path(filepath)

    # 1. TXT 直接读取
    if ext == "txt":
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

    # 2. PDF 用 pypdf
    if ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            text = "\n".join(pages)
            if len(text.strip()) >= 30:
                return text
        except Exception:
            pass
        return ""

    # 3. DOCX 用 zipfile + xml 直接解析（兼容性最好）
    if ext == "docx":
        try:
            import zipfile
            from xml.etree import ElementTree
            with zipfile.ZipFile(filepath) as archive:
                xml = archive.read("word/document.xml")
            root = ElementTree.fromstring(xml)
            paragraphs = []
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            for paragraph in root.iter(ns + "p"):
                runs = [node.text or "" for node in paragraph.iter(ns + "t")]
                text = "".join(runs).strip()
                if text:
                    paragraphs.append(text)
            return "\n".join(paragraphs)
        except Exception:
            pass
        # Fallback: try python-docx
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            pass
        return ""

    # 4. DOC 旧格式：python-docx 不支持，尝试读取原始字节中的可打印文本
    if ext == "doc":
        try:
            # .doc 是 OLE 二进制格式，简单提取可读文本作为 fallback
            with open(filepath, "rb") as f:
                raw = f.read()
            # 尝试用 python-docx（某些 .doc 实际上是 .docx）
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(filepath)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if len(text.strip()) >= 20:
                    return text
            except Exception:
                pass
            # 二进制中提取可打印中英文字符
            import re
            # 解码为 latin-1 保留所有字节
            text = raw.decode("latin-1", errors="ignore")
            # 提取中英文可读片段
            chunks = re.findall(r'[一-鿿　-〿＀-￯a-zA-Z0-9\s.,;:!?()（）、。，；：！？…—\-/+*=#@]{3,}', text)
            result = " ".join(chunks)
            if len(result) >= 20:
                return result
            return ""
        except Exception:
            return ""

    return ""


# ============================================================
# 简历解析
# ============================================================

RESUME_PARSE_PROMPT = """你是一名专业 HR 与职业分析师。请仔细分析以下简历内容，并严格输出 JSON 格式的分析结果。

输出 JSON 结构：
{
  "education": {
    "school": "学校名称",
    "major": "专业",
    "degree": "学历（本科/硕士/博士/大专）",
    "grade": "年级或毕业年份"
  },
  "skills": [
    {"name": "技能名", "category": "编程语言/框架/工具/领域知识/软技能", "level": "入门/熟练/精通/专家", "score": 0-100}
  ],
  "projects": [
    {"name": "项目名", "description": "简要描述", "skills_used": ["技能1", "技能2"]}
  ],
  "advantages": ["优势1", "优势2"],
  "weaknesses": ["不足1", "不足2"],
  "overall_score": 0-100
}

注意：
- skills 数组至少列出 3-8 个技能
- score 分值要合理，不要全部 90+
- category 必须从给定的五个选项中选择
- 只输出 JSON，不要任何其他文字"""


def parse_resume_text(text: str) -> dict:
    """调用 DeepSeek 解析简历文本，返回结构化数据"""
    if not text or not text.strip():
        return {"error": "简历内容为空", "skills": [], "education": {}, "projects": [], "overall_score": 0}

    content, meta = deepseek.chat_completions(
        [
            {"role": "system", "content": RESUME_PARSE_PROMPT},
            {"role": "user", "content": f"简历内容:\n{text[:8000]}"},
        ],
        temperature=0.3,
        timeout=90.0,
    )

    if meta.get("error"):
        return {"error": meta["error"], "skills": [], "education": {}, "projects": [], "overall_score": 0}

    try:
        return deepseek._extract_json(content)
    except Exception:
        return {"error": "AI 返回格式异常，请重试", "raw": content[:500], "skills": [], "education": {}, "projects": [], "overall_score": 0}


# ============================================================
# AI 职业访谈
# ============================================================

INTERVIEW_SYSTEM_PROMPT = """# 智能职业访谈 Agent

## 一、角色定义

你是一名专业的「智能职业访谈与职业画像分析助手」。

你的核心任务不是简单聊天，而是通过自然、循序渐进的职业访谈，帮助用户完成个人职业信息采集。

最终目标是：
1. 获取用户真实、完整、可验证的个人职业信息；
2. 帮助没有现成简历的用户梳理教育经历、技能、项目、实践经历和个人优势；
3. 将非结构化的聊天内容转化为结构化职业档案；
4. 为后续的简历生成、能力画像、岗位匹配和职业发展建议提供数据基础。

必须始终遵循：**先采集事实，再进行分析；先了解用户，再提出建议。**
不得在信息不足时凭空推测用户经历、技能水平或职业能力。

## 二、访谈原则

### 自然对话
不要让用户感觉正在填写机械问卷。像专业职业顾问一样交流。
错误：「请填写您的专业、学历、技能。」
正确：「先从你的学习经历开始吧。你现在学的是什么专业？」

### 动态追问
每次最多询问 1～2 个核心问题。根据用户上一轮回答决定下一轮问题。

## 三、访谈六阶段

阶段1 基础身份：姓名、学历、学校、专业、年级/毕业年份、当前身份
阶段2 教育经历：专业方向、核心课程、感兴趣方向、获奖、证书
阶段3 技能能力：编程语言、工具、框架、数据库、AI/ML、办公工具。每项技能追问使用场景和熟练度，不要单凭用户说"会"就判定熟练
阶段4 项目与实践（最重要）：课程项目、实训、比赛、科研、实习、社团、个人项目、开源。对每个项目追问：做什么/为什么/你负责什么/用了什么技术/遇到什么问题/怎么解决/结果/量化成果
阶段5 职业目标：目标岗位、感兴趣行业、期望城市、求职状态、发展方向
阶段6 个人优势与偏好：自评优势、擅长问题、工作偏好、学习/沟通/协作能力

## 四、信息可信度规则

将用户信息分为三类：
- confirmed：用户明确说明且信息清晰
- self_assessed：用户主观评价（如"我觉得自己Python比较熟练"）
- inferred：根据经历推断（必须明确标记，不可伪装成用户明确提供的信息）

## 五、核心规则

1. 禁止编造用户没有提供的经历、项目成果、实习经历
2. 禁止擅自提高用户技能等级
3. 禁止将推测当成事实
4. 禁止一次提出大量问题
5. 禁止重复询问已确认的信息
6. **重要**：当用户明确要求结束访谈/生成结果时（如"请结束""生成档案""就这样吧""已经说完了""请基于我们讨论的内容结束"），必须立即设置 is_complete: true 并输出完整的 profile_data。不要再追问。缺失信息标记为 missing_information 即可。
7. 如果用户回答"不知道"，不要强迫，可以提供选项引导
8. 没有项目经历的用户，进一步挖掘课程作业、实验、个人练习、比赛等
9. 回答太简单的用户，继续追问使用场景
10. 如果一轮对话中用户一次性提供了大量信息（多个阶段的答案），应当先确认理解，然后判断是否已经满足完成条件。不要忽略用户已提供的信息继续追问

## 六、对话风格

专业、自然、简洁、鼓励、不审问、不说教、有职业顾问感。
每次回复 2～5 句话，不要长篇解释。
让用户感觉"有人在帮我一步步梳理经历"，而不是"我在填表"。

## 七、输出规则

访谈进行中：只输出自然语言对话，不展示JSON、不展示内部字段、不说"正在提取技能标签"。

每轮必须输出 JSON：
{"reply": "你对用户说的话（自然、专业、简洁）", "is_complete": false, "questions_asked": 1, "profile_data": null}

访谈完成条件（必须信息：基础身份 + 教育经历 + 至少一项技能 + 职业方向）满足后：
{"reply": "总结话语", "is_complete": true, "questions_asked": N, "profile_data": {完整结构化档案}}

profile_data 格式：
{
  "basic_info": {"name":"", "education":"", "school":"", "major":"", "grade":"", "graduation_year":""},
  "education": {"courses":[], "achievements":[], "certificates":[]},
  "skills": [{"name":"", "level":"", "evidence":"", "confidence":"confirmed|self_assessed|inferred"}],
  "projects": [{"name":"", "description":"", "role":"", "technologies":[], "responsibilities":[], "results":"", "confidence":"confirmed|self_assessed|inferred"}],
  "experience": [],
  "career_goal": {"target_jobs":[], "industries":[], "preferred_city":"", "career_direction":""},
  "personal_strengths": [],
  "career_preferences": [],
  "missing_information": [],
  "profile_summary": "",
  "interview_completed": true
}"""


def run_career_interview(history: list[dict], user_message: str) -> dict:
    """运行 AI 职业访谈对话"""
    messages = [{"role": "system", "content": INTERVIEW_SYSTEM_PROMPT}]
    messages.extend(history[-20:])  # 保留最近 20 轮对话
    messages.append({"role": "user", "content": user_message})

    content, meta = deepseek.chat_completions(
        messages,
        temperature=0.4,
        timeout=90.0,
    )

    if meta.get("error"):
        return {"reply": "抱歉，AI 服务暂时不可用，请稍后重试。", "is_complete": False, "profile_data": None, "error": meta["error"]}

    try:
        return deepseek._extract_json(content)
    except Exception:
        return {"reply": content[:500], "is_complete": False, "profile_data": None}


def start_interview() -> dict:
    """开始一次新的访谈，返回开场白"""
    content, meta = deepseek.chat_completions(
        [
            {"role": "system", "content": INTERVIEW_SYSTEM_PROMPT},
            {"role": "user", "content": "你好，请开始职业访谈。"},
        ],
        temperature=0.4,
        timeout=90.0,
    )
    if meta.get("error"):
        return {"reply": "你好！我是你的职业规划助手。请先告诉我你的专业背景吧～", "is_complete": False, "profile_data": None}
    try:
        return deepseek._extract_json(content)
    except Exception:
        return {"reply": "你好！我是你的职业规划助手。让我们开始吧——请告诉我你的专业背景？", "is_complete": False, "profile_data": None}


# ============================================================
# 从访谈数据生成简历
# ============================================================

RESUME_GEN_PROMPT = """你是一位专业简历撰写顾问。根据以下职业访谈得到的结构化数据，为用户生成一份正式的中文简历。

要求：
1. 使用标准的简历格式：个人信息、教育背景、技能特长、项目经历、实践经历、自我评价
2. 只使用数据中标记为 confirmed 和 self_assessed 的信息
3. inferred 信息仅作参考，不直接写入简历
4. 语言简洁专业，量化成果优先
5. 技能按熟练度排序，标注熟练程度
6. 不要编造任何数据中没有的信息

输出纯文本简历（Markdown 格式），不要 JSON。"""


def generate_resume_from_profile(profile_data: dict) -> str:
    """根据访谈画像数据，用 AI 生成简历"""
    # 构建精简的 prompt context
    basic = profile_data.get("basic_info", {})
    skills = [s for s in profile_data.get("skills", []) if s.get("confidence") in ("confirmed", "self_assessed")]
    projects = [p for p in profile_data.get("projects", []) if p.get("confidence") in ("confirmed", "self_assessed")]
    career = profile_data.get("career_goal", {})

    context = {
        "basic": basic,
        "education": profile_data.get("education", {}),
        "skills": skills,
        "projects": projects,
        "experience": profile_data.get("experience", []),
        "career_goal": career,
        "strengths": profile_data.get("personal_strengths", []),
        "summary": profile_data.get("profile_summary", ""),
    }

    content, meta = deepseek.chat_completions(
        [
            {"role": "system", "content": RESUME_GEN_PROMPT},
            {"role": "user", "content": f"职业档案数据:\n{json.dumps(context, ensure_ascii=False)[:6000]}"},
        ],
        temperature=0.3,
        timeout=60.0,
    )

    if meta.get("error"):
        return _build_simple_resume(profile_data)

    return content.strip() or _build_simple_resume(profile_data)


def _build_simple_resume(profile_data: dict) -> str:
    """无 AI 时的简单简历拼装"""
    basic = profile_data.get("basic_info", {})
    skills = profile_data.get("skills", [])
    projects = profile_data.get("projects", [])
    career = profile_data.get("career_goal", {})
    edu = profile_data.get("education", {})
    strengths = profile_data.get("personal_strengths", [])

    lines = [
        f"# {basic.get('name', '未填写')}",
        f"",
        f"## 个人信息",
        f"- 学校：{basic.get('school', '--')}",
        f"- 专业：{basic.get('major', '--')}",
        f"- 学历：{basic.get('education', '--')}",
        f"- 年级：{basic.get('grade', '--')}",
        f"- 目标岗位：{', '.join(career.get('target_jobs', [])) or '--'}",
        f"",
        f"## 教育背景",
    ]
    for c in edu.get("courses", []):
        lines.append(f"- {c}")
    for a in edu.get("achievements", []):
        lines.append(f"- 🏆 {a}")
    for c in edu.get("certificates", []):
        lines.append(f"- 📜 {c}")

    lines.append(f"")
    lines.append(f"## 技能特长")
    for s in skills:
        conf = "✅" if s.get("confidence") == "confirmed" else "📝"
        lines.append(f"- {conf} {s.get('name', '')} — {s.get('level', '熟练')}")

    lines.append(f"")
    lines.append(f"## 项目经历")
    for p in projects:
        lines.append(f"### {p.get('name', '项目')}")
        lines.append(f"- 角色：{p.get('role', '--')}")
        lines.append(f"- 描述：{p.get('description', '--')}")
        lines.append(f"- 技术：{', '.join(p.get('technologies', []))}")
        lines.append(f"- 成果：{p.get('results', '--')}")

    if strengths:
        lines.append(f"")
        lines.append(f"## 自我评价")
        for s in strengths:
            lines.append(f"- {s}")

    lines.append(f"")
    lines.append(f"*本简历由 AI 职业访谈自动生成*")
    return "\n".join(lines)


# ============================================================
# 岗位匹配
# ============================================================

def calculate_job_match(user_skills: list[dict], top_n: int = 10) -> list[dict]:
    """计算用户技能与 job_postings 岗位的匹配度"""
    db = SessionLocal()
    try:
        from backend.db_models import JobPosting, JobPostingDetail

        # 取最近 200 条有效岗位（过滤掉空标题）
        jobs = (
            db.query(JobPosting)
            .filter(JobPosting.job_title.isnot(None), JobPosting.job_title != '')
            .order_by(JobPosting.crawl_time.desc().nullslast())
            .limit(300)
            .all()
        )

        user_skill_names = set(s.get("name", "").lower() for s in user_skills)
        user_skill_map = {s.get("name", "").lower(): s.get("score", 50) for s in user_skills}

        results = []
        for job in jobs:
            jd = (
                db.query(JobPostingDetail)
                .filter(JobPostingDetail.job_id == job.id)
                .first()
            )

            # 提取职位技能关键词
            job_skills = set()
            if jd and jd.skills:
                job_skills = set(s.lower() for s in (jd.skills or []))
            if jd and jd.keywords:
                job_skills |= set(k.lower() for k in (jd.keywords or []))

            if not job_skills:
                # 从职位标题提取关键词作为退路
                title_words = set(w.lower() for w in re.findall(r"[\w一-鿿]+", job.job_title or ""))
                job_skills = title_words

            # 技能匹配度 (50%)
            if job_skills:
                matched = user_skill_names & job_skills
                skill_match = len(matched) / max(len(job_skills), 1)
            else:
                skill_match = 0.5  # 无技能数据时给中值
                matched = set()

            # 学历匹配 (20%)
            edu_match = 0.7  # 默认中高，因为大多数岗位学历要求宽泛
            if jd and jd.education_required:
                edu_match = 0.8  # 宽松匹配

            # 经验匹配 (20%)
            exp_match = 0.5  # 默认中等

            # 方向匹配 (10%)
            direction_match = 0.5

            # 综合分数
            total = (
                skill_match * 0.50
                + edu_match * 0.20
                + exp_match * 0.20
                + direction_match * 0.10
            )
            match_score = round(total * 100)

            results.append(
                {
                    "job_id": job.id,
                    "title": job.job_title,
                    "company": job.company_name,
                    "city": job.city or "",
                    "salary_min": job.salary_min,
                    "salary_max": job.salary_max,
                    "education": job.education or "",
                    "experience": job.experience or "",
                    "match_score": match_score,
                    "matched_skills": list(matched)[:5],
                }
            )

        results.sort(key=lambda x: x["match_score"], reverse=True)
        return results[:top_n]

    finally:
        db.close()


# ============================================================
# 画像构建
# ============================================================

def build_radar_data(skills: list[dict]) -> dict:
    """将技能列表映射为 6 维雷达图数据"""
    # 按类别聚合评分
    dims = {
        "tech": {"name": "技术能力", "scores": []},
        "project": {"name": "项目经验", "scores": []},
        "data": {"name": "数据能力", "scores": []},
        "engineering": {"name": "工程能力", "scores": []},
        "innovation": {"name": "创新能力", "scores": []},
        "learning": {"name": "学习能力", "scores": []},
    }

    category_to_dim = {
        "编程语言": "tech",
        "框架": "engineering",
        "工具": "engineering",
        "领域知识": "data",
        "软技能": "learning",
    }

    for s in (skills or []):
        cat = s.get("category", "")
        dim_key = category_to_dim.get(cat, "tech")
        if dim_key in dims:
            dims[dim_key]["scores"].append(s.get("score", 50))

    indicators = []
    values = []
    for key, d in dims.items():
        avg = round(sum(d["scores"]) / max(len(d["scores"]), 1)) if d["scores"] else 30
        indicators.append({"name": d["name"], "max": 100})
        values.append(avg)

    return {
        "indicators": indicators,
        "series": [{"name": "能力画像", "value": values}],
    }


def calculate_overall(skills: list[dict], radar_values: list[int]) -> int:
    """计算综合职业竞争力评分"""
    skill_avg = 0
    if skills:
        skill_avg = sum(s.get("score", 50) for s in skills) / len(skills)
    radar_avg = sum(radar_values) / max(len(radar_values), 1) if radar_values else 0
    return round(skill_avg * 0.6 + radar_avg * 0.4)


def save_career_report(
    user_id: str,
    skills: list[dict],
    radar_data: dict,
    advantages: list[str],
    weaknesses: list[str],
    suggestions: list[str],
    match_jobs: list[dict],
    raw_analysis: dict | None = None,
) -> CareerReport:
    """保存职业分析报告到数据库"""
    db = SessionLocal()
    try:
        dims = {d["name"]: v for d, v in zip(radar_data.get("indicators", []), radar_data.get("series", [{}])[0].get("value", []))}
        overall = calculate_overall(skills, list(dims.values()))

        report = CareerReport(
            user_id=user_id,
            tech_score=dims.get("技术能力", 0),
            project_score=dims.get("项目经验", 0),
            created_at=datetime.now(),
            data_score=dims.get("数据能力", 0),
            engineering_score=dims.get("工程能力", 0),
            innovation_score=dims.get("创新能力", 0),
            learning_score=dims.get("学习能力", 0),
            overall_score=overall,
            advantages=advantages or [],
            weaknesses=weaknesses or [],
            suggestions=suggestions or [],
            match_jobs=match_jobs or [],
            raw_analysis=raw_analysis or {},
        )
        db.add(report)
        db.commit()
        db.refresh(report)
        return report
    finally:
        db.close()


def save_skills_from_parse(user_id: str, skills: list[dict], source: str = "resume") -> list[UserSkill]:
    """将解析出的技能保存到数据库"""
    db = SessionLocal()
    try:
        # 删除旧数据（同一来源）
        db.query(UserSkill).filter(UserSkill.user_id == user_id, UserSkill.source == source).delete()

        saved = []
        for s in (skills or []):
            skill = UserSkill(
                user_id=user_id,
                skill_name=s.get("name", ""),
                category=s.get("category", "领域知识"),
                level=s.get("level", "熟练"),
                score=s.get("score", 50),
                source=source,
            )
            db.add(skill)
            saved.append(skill)

        db.commit()
        return saved
    finally:
        db.close()


# ============================================================
# AI 简历优化
# ============================================================

OPTIMIZE_PROMPT = """你是一名专业简历优化顾问。根据原始简历内容 + AI分析结果，对简历进行优化。

核心原则：FACT PRESERVATION — 只能重组/润色/提炼已有的真实信息，严禁编造任何用户未提供的内容。

优化规则：
1. 保留所有真实事实，不虚构任何信息（工作经历、项目、公司、数据、成果等）
2. 优化语言表达：提升职业化程度、技术表达准确性
3. 优化项目描述：动作 + 技术 + 任务（原始有成果才写成果）
4. 优化技能表达：突出已有技能的实际应用场景
5. 删除重复、无效、冗余表达
6. 原始信息不足时保持为空，不自行补充
7. 禁止生成虚假量化指标（如"处理100万条数据""提升35%"等）

根据优化模式调整强度：
- light(轻度)：只优化语言表达、错别字、语句通顺度，尽量保留原文
- professional(专业)：优化项目描述、技能表达、职业化程度、关键词、结构
- deep(深度)：整体重构（个人简介/技能/项目/经历描述），但不改变事实

输出 JSON 格式（原始内容使用原文，优化后使用优化文本）：
{
  "summary": {"original": "原始个人简介或概述（无则为空字符串）", "optimized": "优化后"},
  "skills": [{"original": "原始技能描述", "optimized": "优化后技能表达", "reason": "优化原因"}],
  "projects": [{"name": "项目名", "original": "原始项目描述", "optimized": "优化后", "reason": "优化原因"}],
  "highlights": [{"type": "expression|structure|keyword|professionalism", "original": "原文片段", "optimized": "优化后", "reason": "原因"}],
  "optimization_summary": ["整体优化说明1", "整体优化说明2"],
  "changes_count": 0
}

只输出 JSON，不要其他文字。"""


def optimize_resume(original_text: str, ai_analysis: dict, mode: str = "professional") -> dict:
    """调用 DeepSeek 优化简历，返回结构化对比数据"""
    if not original_text or not original_text.strip():
        return {"error": "原始简历内容为空"}

    mode_labels = {"light": "轻度优化", "professional": "专业优化", "deep": "深度优化"}
    mode_label = mode_labels.get(mode, "专业优化")

    # 精简 AI 分析结果
    compact_analysis = {
        "education": ai_analysis.get("education", {}),
        "skills": [s.get("name", "") for s in ai_analysis.get("skills", [])[:10]],
        "projects": [p.get("name", "") for p in ai_analysis.get("projects", [])[:5]],
        "advantages": ai_analysis.get("advantages", [])[:5],
        "weaknesses": ai_analysis.get("weaknesses", [])[:5],
        "overall_score": ai_analysis.get("overall_score", 0),
    }

    content, meta = deepseek.chat_completions(
        [
            {"role": "system", "content": OPTIMIZE_PROMPT},
            {"role": "user", "content": f"优化模式: {mode_label}\n\n原始简历:\n{original_text[:6000]}\n\nAI分析结果:\n{json.dumps(compact_analysis, ensure_ascii=False)}"},
        ],
        temperature=0.3,
        timeout=120.0,
    )

    if meta.get("error"):
        return {"error": meta["error"], "fallback": True}

    try:
        result = deepseek._extract_json(content)
        result["mode"] = mode
        result["ai_error"] = None
        return result
    except Exception:
        return {"error": "AI 返回格式异常", "raw": content[:500], "fallback": True}
        db.close()


# ============================================================
# AI 访谈分析
# ============================================================

INTERVIEW_ANALYSIS_PROMPT = """你是一名专业的职业分析师。根据用户与AI职业顾问的对话内容，生成一份职业信息分析报告。

规则：
1. 只分析访谈中用户明确表达的信息
2. 不编造用户没有提到的技能、经历、项目、证书
3. 不编造用户没有表达的职业兴趣
4. 信息不足时标记为"暂未提及"而不是猜测

输出 JSON：
{
  "overview": {"stage": "职业发展阶段", "direction": "兴趣方向", "goal": "职业目标", "summary": "一句话概述"},
  "core_abilities": [{"name": "能力名", "description": "来自访谈的证据"}],
  "career_interests": [{"name": "兴趣方向", "evidence": "用户表达"}],
  "work_preferences": [{"name": "偏好", "description": "说明"}],
  "strengths": [{"title": "优势", "detail": "来自访谈"}],
  "improvements": [{"title": "方向", "suggestion": "建议"}],
  "career_directions": [{"title": "推荐岗位", "reason": "基于访谈的推荐理由"}],
  "summary": "基于本次访谈的职业总结"
}

只输出JSON。"""


def analyze_interview(conversation: list[dict]) -> dict:
    """分析访谈对话内容，生成职业分析报告"""
    if not conversation:
        return {"error": "访谈内容为空"}

    # 构建精简的对话上下文
    lines = []
    for m in conversation[-20:]:  # 最近20轮
        role = "用户" if m.get("role") in ("user",) else "AI顾问"
        text = m.get("text", m.get("content", ""))
        if text:
            lines.append(f"{role}: {text[:300]}")
    dialog = "\n".join(lines)

    content, meta = deepseek.chat_completions(
        [
            {"role": "system", "content": INTERVIEW_ANALYSIS_PROMPT},
            {"role": "user", "content": f"职业访谈对话:\n{dialog[:6000]}"},
        ],
        temperature=0.3,
        timeout=90.0,
    )

    if meta.get("error"):
        return {"error": meta["error"]}

    try:
        return deepseek._extract_json(content)
    except Exception:
        return {"error": "AI 返回格式异常", "raw": content[:500]}
