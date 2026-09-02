# -*- coding: utf-8 -*-
"""Inspect reference competition plan DOCX structure."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

REF = Path(r"N:\ppt素材\多份国赛获奖作品PPT\【国赛】石墨烯基柔性锂离子电容器-柔性储能先锋（附PPT+视频）-\石墨烯基柔性锂离子电容器-柔性储能先锋_计划书---淘宝店：-馆.docx")
CUR = Path(r"C:\Users\Ibiza\Desktop\project\挑战杯\作品设计实现方案_提交版.docx")
OUT = Path(r"C:\Users\Ibiza\Desktop\project\挑战杯\scripts\devtools\_ref_plan_inspect.txt")


def align_name(p):
    try:
        a = p.alignment
        if a is None and p.paragraph_format.alignment is not None:
            a = p.paragraph_format.alignment
        if a is None:
            return "None"
        return str(a).split(".")[-1]
    except Exception:
        return "?"


def dump(doc_path, label, limit=120):
    doc = Document(str(doc_path))
    lines = []
    lines.append(f"===== {label} =====")
    lines.append(f"path: {doc_path}")
    lines.append(f"paragraphs: {len(doc.paragraphs)}")
    lines.append(f"tables: {len(doc.tables)}")
    lines.append(f"sections: {len(doc.sections)}")
    for i, sec in enumerate(doc.sections):
        lines.append(
            f"  sec{i}: page={sec.page_width.cm:.1f}x{sec.page_height.cm:.1f}cm "
            f"margins LRTB={sec.left_margin.cm:.2f}/{sec.right_margin.cm:.2f}/{sec.top_margin.cm:.2f}/{sec.bottom_margin.cm:.2f}"
        )
    # styles used
    styles = {}
    for p in doc.paragraphs:
        sn = p.style.name if p.style else "?"
        styles[sn] = styles.get(sn, 0) + 1
    lines.append("styles: " + ", ".join(f"{k}:{v}" for k, v in sorted(styles.items(), key=lambda x: -x[1])[:20]))

    # images
    rels = doc.part.rels
    imgs = [r for r in rels.values() if "image" in r.reltype]
    lines.append(f"images: {len(imgs)}")

    lines.append("--- first paragraphs ---")
    shown = 0
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if not t and not p.runs:
            # skip empties after first few
            if shown > 40:
                continue
        style = p.style.name if p.style else "?"
        fonts = []
        for r in p.runs[:3]:
            fonts.append(f"{r.font.name}/{r.font.size.pt if r.font.size else '-'}{'b' if r.font.bold else ''}")
        lines.append(f"[{i}] ({style}|{align_name(p)}) {t[:120]!r}  runs={fonts}")
        shown += 1
        if shown >= limit:
            break

    # headings outline
    lines.append("--- headings ---")
    for i, p in enumerate(doc.paragraphs):
        sn = p.style.name if p.style else ""
        if sn.startswith("Heading") or sn.startswith("标题") or "标题" in sn:
            lines.append(f"  {sn}: {(p.text or '')[:80]}")
    return "\n".join(lines)


text = dump(REF, "REFERENCE", 150) + "\n\n" + dump(CUR, "CURRENT", 80)
OUT.write_text(text, encoding="utf-8")
print(text[:8000])
print("\n... wrote", OUT)
