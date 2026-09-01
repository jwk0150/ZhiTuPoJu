# -*- coding: utf-8 -*-
"""按专家批注修订 lunwen 副本，保存到桌面。"""
from __future__ import annotations

import copy
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from lxml import etree

SRC = Path(r"C:\Users\Ibiza\Desktop\lunwen - 副本.docx")
DST = Path(r"C:\Users\Ibiza\Desktop\lunwen-按批注修订.docx")


def set_run_east_asia(run, font_name: str, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def clear_paragraph(p):
    for child in list(p._element):
        # keep pPr
        if child.tag == qn("w:pPr"):
            continue
        p._element.remove(child)


def set_paragraph_text(p, text: str, font_name="宋体", size_pt=12, bold=False, first_line_indent=True):
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_east_asia(run, font_name, size_pt, bold)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)
    else:
        pf.first_line_indent = None


def set_heading_text(p, text: str, level: int):
    """Rewrite heading text and force 黑体."""
    clear_paragraph(p)
    run = p.add_run(text)
    # H1 小二/三号约 16pt，H2 四号 14pt，H3 小四 12pt — 与常见中文论文习惯接近
    size = {1: 16, 2: 14, 3: 12}.get(level, 12)
    set_run_east_asia(run, "黑体", size, True)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def clone_paragraph_after(ref_p, src_p):
    """Insert a deep copy of src_p after ref_p, return new paragraph element wrapper-ish."""
    new_el = copy.deepcopy(src_p._element)
    ref_p._element.addnext(new_el)
    return new_el


def delete_paragraph(p):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def strip_label_colon(text: str) -> str:
    """去掉『标题：』式起笔，保留后文。"""
    # fullwidth and halfwidth colon
    m = re.match(r"^[^：:]{1,20}[：:]\s*(.*)$", text.strip())
    if m and m.group(1):
        return m.group(1)
    return text


def configure_heading2_style(doc: Document):
    h2 = doc.styles["Heading 2"]
    h2.font.name = "黑体"
    h2.font.bold = True
    h2.font.size = Pt(14)
    rPr = h2.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "黑体")


def remove_all_comments(docx_path: Path):
    """Remove comment parts so revised file is clean."""
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            name = item.filename
            data = zin.read(name)
            if name.startswith("word/comments") or name.startswith("word/comments"):
                continue
            if name == "word/document.xml":
                root = etree.fromstring(data)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                # remove comment range markers / refs
                for tag in (
                    "commentRangeStart",
                    "commentRangeEnd",
                    "commentReference",
                ):
                    for el in root.findall(f".//w:{tag}", ns):
                        parent = el.getparent()
                        if parent is not None:
                            parent.remove(el)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            if name == "[Content_Types].xml":
                root = etree.fromstring(data)
                for el in list(root):
                    pn = el.get("PartName", "")
                    if "comments" in pn.lower():
                        root.remove(el)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            if name == "word/_rels/document.xml.rels":
                root = etree.fromstring(data)
                for el in list(root):
                    target = el.get("Target", "")
                    if "comment" in target.lower():
                        root.remove(el)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    tmp.replace(docx_path)


def main():
    doc = Document(str(SRC))
    configure_heading2_style(doc)
    paras = list(doc.paragraphs)

    # ---------- 摘要：去掉“消融”硬写 ----------
    abstract = paras[4].text
    abstract = abstract.replace("全文对公式、伪代码、对照实验与消融作统一整理，使创新主张可被数据与实现共同支撑。",
                                "全文对公式、伪代码与对照实验作统一整理，使创新主张可被数据与实现共同支撑。")
    set_paragraph_text(paras[4], abstract, bold=True, first_line_indent=True)
    # 保持摘要加粗观感：上面已 bold=True

    abs_en = paras[10].text
    abs_en = abs_en.replace("baselines,and ablations are reported so that claims remain checkable.",
                            "baselines are reported so that claims remain checkable.")
    abs_en = abs_en.replace("baselines, and ablations are reported so that claims remain checkable.",
                            "baselines are reported so that claims remain checkable.")
    set_paragraph_text(paras[10], abs_en, font_name="Times New Roman", size_pt=10.5, bold=True, first_line_indent=True)

    # ---------- 引言：问题先行，再引出平台 ----------
    intro = (
        "当前就业市场面临技术栈迭代快于人才培养周期、企业“招人难”与劳动者“路径不清”并存的结构性矛盾。"
        "公开招聘文本虽是观测能力需求的重要信号源，却普遍存在发布时滞、模板化复制与表述噪声，"
        "直接污染下游知识图谱构建与人岗匹配结果；与此同时，传统招聘多依赖关键词匹配，"
        "难以刻画岗位能力结构的动态演化，而大模型在岗位定义生成与诊断对话中又易出现证据不足的幻觉。"
        "针对上述问题，本文设计并实现面向就业市场的岗位—能力知识图谱动态构建与智能匹配系统「执图破局」。"
        "系统以岗位与能力为核心实体，融合多源数据采集治理、图谱构建、新兴岗位发现、能力演化差分、"
        "五维人岗匹配与检索增强对话，形成可核验、可复算、可执行的闭环。"
        "在线评测中，JD解析、简历提取与匹配准确率均达90%以上，可支撑企业招聘、职业规划与人才态势研判。"
    )
    set_paragraph_text(paras[13], intro)

    # ---------- 相关工作：去掉“国外研究现状：”标签体 ----------
    set_paragraph_text(
        paras[15],
        "国外在岗位—技能知识图谱领域研究相对成熟，已形成较完善的职业分类、技能体系和岗位标准，"
        "并进一步融合招聘信息、课程资源等动态数据构建岗位—技能知识图谱。"
        "例如相关研究利用岗位招聘数据不断更新技能与职业之间的关联，并结合图神经网络、知识图谱推理和大语言模型"
        "开展技能预测、岗位匹配与职业路径规划。近年来，研究重点逐渐由静态知识表示转向时序建模和动态演化，"
        "以适应劳动力市场快速变化的需求。",
    )
    set_paragraph_text(
        paras[16],
        "国内关于岗位—能力知识图谱的研究起步较晚，目前主要集中于招聘信息挖掘、职业能力分析和人才岗位匹配等方向。"
        "研究者逐渐利用招聘网站、职业标准、课程资源等多源异构数据，通过自然语言处理和知识图谱技术提取岗位、技能及其关联关系，"
        "实现岗位需求的结构化表达，并进一步应用于人才推荐和职业规划。"
        "但现有研究仍以静态知识组织和单一场景应用为主，对岗位需求随产业发展产生的动态变化及知识图谱持续演化研究相对不足。",
    )
    set_paragraph_text(
        paras[17],
        "总体来看，国内外研究均已认识到多源数据融合和知识图谱在岗位能力分析中的重要价值，"
        "但仍存在数据来源分散、异构数据融合困难、岗位与能力关系动态变化刻画不足等问题。"
        "现有方法较多关注知识图谱构建或岗位匹配的单一环节，对“数据采集—知识抽取—图谱构建—动态更新—演化分析”的完整闭环研究仍不充分。"
        "因此，有必要构建面向就业市场的多源异构岗位—能力知识图谱，实现岗位需求变化和能力结构演化的动态刻画。",
    )

    # ---------- 第三章：按 3.1结构 / 3.2创新 / 3.3实现 重组 ----------
    # 先改各级标题文本与字体
    set_heading_text(paras[18], "三、系统设计与实现", 1)
    set_heading_text(paras[19], "3.1 系统结构", 2)

    # 删除 Input-Hidden-Output 框架段与图1说明、设计思路（批注：去掉）
    # 20 DRAW(含图1), 21 图题, 22 设计思路
    for idx in (22, 21, 20):
        delete_paragraph(paras[idx])

    # 重新取段落列表（删除后索引变化）——改用稳定策略：按文本匹配后续操作
    paras = list(doc.paragraphs)

    def find_para(startswith: str | None = None, equals: str | None = None, contains: str | None = None):
        for p in doc.paragraphs:
            t = p.text.strip()
            if equals is not None and t == equals:
                return p
            if startswith is not None and t.startswith(startswith):
                return p
            if contains is not None and contains in t:
                return p
        return None

    # 3.1 开头插入系统结构总述（在原“据质量自治模型”之前）
    p_data = find_para(startswith="据质量自治模型") or find_para(contains="岗位主数据拆成查询总表")
    assert p_data is not None

    # 在 p_data 前插入结构总述
    new_el = OxmlElement("w:p")
    p_data._element.addprevious(new_el)
    # wrap as paragraph via body
    from docx.text.paragraph import Paragraph
    p_struct = Paragraph(new_el, p_data._parent)
    set_paragraph_text(
        p_struct,
        "系统总体分为数据治理层、知识图谱层、智能分析层与应用交互层。"
        "数据治理层完成多源招聘与简历语料的接入、指纹去重、完整度评分与交叉验证；"
        "知识图谱层将岗位、技能、行业、公司及其关系组织为可查询的有向加权图；"
        "智能分析层承担新兴岗位发现、既有岗位能力演化差分与五维人岗匹配；"
        "应用交互层输出结构化岗位定义、演化报告、匹配诊断与受证据约束的对话答复。"
        "设计约束是：任何写入“新兴岗位定义”或“技能已删除”的结论，必须能指回具体样本集合或可复算公式，"
        "从而保证入库数据可核验、图谱关系可查询、演化差分可复算、匹配结论可执行。",
    )

    # 改写数据层段落（去掉“据质量自治模型：”标签体）
    data_text = p_data.text
    data_text = strip_label_colon(data_text)
    # 原文开头缺字“数”，补上
    if data_text.startswith("岗位主数据"):
        data_text = (
            "在数据层，岗位主数据拆成查询总表与细节表，二者通过外键1:1关联，如图2所示。"
            + data_text[len("岗位主数据拆成查询总表与细节表，二者通过外键1:1关联，如图2所示。"):]
            if "岗位主数据拆成查询总表与细节表，二者通过外键1:1关联，如图2所示。" in data_text
            else "在数据层，" + data_text
        )
    # 更稳妥的重写
    set_paragraph_text(
        p_data,
        "在数据层，岗位主数据拆成查询总表与细节表，二者通过外键1:1关联，如图2所示。"
        "总表承载高频筛选字段（来源、标题、公司、城市、薪资、经验、学历、发布时间、指纹、完整度等），"
        "细节表承载长文本与数组型字段（描述、要求、技能数组、福利、分级类目等），"
        "平台特有字段进入JSONB，避免频繁改模式。"
        "总表与细节表满足外键约束，保证查询与详情读取路径一致。",
    )

    # ---------- 把原“创新点I1…I5”收到 3.2；原 3.2–3.10 收到 3.3 ----------
    # 找到创新点段落与原3.2标题
    p_i1 = find_para(startswith="创新点I1")
    p_old32 = find_para(equals="3.2图谱形式化定义") or find_para(startswith="3.2")
    assert p_i1 is not None and p_old32 is not None

    # 在 I1 前插入 3.2 标题
    h32 = OxmlElement("w:p")
    p_i1._element.addprevious(h32)
    p_h32 = Paragraph(h32, p_i1._parent)
    p_h32.style = doc.styles["Heading 2"]
    set_heading_text(p_h32, "3.2 系统创新", 2)

    # 插入创新总起段
    intro_i = OxmlElement("w:p")
    p_i1._element.addprevious(intro_i)
    p_intro_i = Paragraph(intro_i, p_i1._parent)
    set_paragraph_text(
        p_intro_i,
        "针对数据噪声、动态演化不可复算、匹配虚高与生成幻觉等关键痛点，系统形成五项可复核创新。"
        "各项创新均给出形式化定义、实现路径与定量证据，避免只陈述概念而无公式支撑。",
    )

    # 改写 I1–I5 标签体为完整段落
    rewrites = {
        "创新点I1": (
            "第一项创新是入库质量门控。对每条岗位记录计算内容指纹与十五字段完整度，"
            "在数据库触发器内自动评分并去重，从源头抑制时滞样本、空壳样本与模板复制进入图谱。形式化为"
        ),
        "创新点I2": (
            "第二项创新是三维新兴度评分与六阶段推理链。对标题簇定义"
        ),
        "创新点I3": (
            "第三项创新是多源伪时序技能差分。用偏早源与偏新源的出现率差判定新增、删除与修改，"
            "把“能力演化”写成可复算差分而非叙事描述。"
            "数据支撑方面，演化与专家标注一致性Kappa为0.76；"
            "示意差分如RAG、Prompt、Spark等可在同一抽取器、同一岗位类下复算。"
        ),
        "创新点I4": (
            "第四项创新是五维加权匹配与迁移上界。总分"
        ),
        "创新点I5": (
            "第五项创新是证据门控与幻觉防控。发现与生成结论要求证据源数量达到设定阈值，"
            "否则标记为low_evidence并强制人工复核；对话侧禁止编造上下文外岗位ID，"
            "并在大模型不可用时回退本地启发式。"
            "数据支撑显示，未审计幻觉率约18.7%，门控后约3.2%；核心匹配与发现路径在无密钥条件下仍可复现。"
        ),
    }
    for p in list(doc.paragraphs):
        t = p.text.strip()
        for key, new_prefix in rewrites.items():
            if t.startswith(key):
                # I3/I5 整段可直接替换；I1/I2/I4 后接公式，只替换引导语
                if key in ("创新点I3", "创新点I5"):
                    set_paragraph_text(p, new_prefix if key == "创新点I5" else new_prefix)
                else:
                    # 保留公式后的“数据支撑”等后半？原段在公式前结束，公式独立段
                    set_paragraph_text(p, new_prefix, first_line_indent=True)
                break

    # I1 后的解释段、I2 后的解释段、I4 后的解释段里去掉“消融”措辞
    for p in doc.paragraphs:
        t = p.text
        if "消融去掉技能维" in t:
            set_paragraph_text(
                p,
                t.replace(
                    "；消融去掉技能维降8.4个百分点，说明权重配置与经验一致。",
                    "。技能维权重最高，与其在匹配中的主导作用相一致。",
                ),
            )

    # 原 3.2 标题改为 3.3，并删除其后多余的 3.3–3.10 二级标题（并入实现叙述）
    p_old32 = find_para(equals="3.2图谱形式化定义") or find_para(contains="图谱形式化定义")
    assert p_old32 is not None
    p_old32.style = doc.styles["Heading 2"]
    set_heading_text(p_old32, "3.3 系统实现", 2)

    # 在 3.3 标题后加一句总起
    impl_intro_el = OxmlElement("w:p")
    p_old32._element.addnext(impl_intro_el)
    p_impl_intro = Paragraph(impl_intro_el, p_old32._parent)
    set_paragraph_text(
        p_impl_intro,
        "本节按“图谱构建—新兴发现—能力演化—人岗匹配—对话防控—实验验证”的顺序给出关键实现。"
        "实现细节与前文创新一一对应，公式与页面结果可相互回证。",
    )

    # 删除旧的 3.3–3.10 标题，改为正文小标题句首（自然段），避免堆叠过多二级标题
    remove_h2_exact = {
        "3.3实体抽取与规范化",
        "3.4关系推理与边权更新",
        "3.5新兴岗位发现与预测",
        "3.6既有岗位能力动态演化",
        "3.7人岗匹配与学习路径",
        "3.8大模型对话与幻觉防控",
        "3.9实验与结果",
        "3.10总结与展望",
    }
    # 也可能带空格
    for p in list(doc.paragraphs):
        t = p.text.strip().replace(" ", "")
        for key in list(remove_h2_exact):
            if t.replace(" ", "") == key.replace(" ", "") or t.startswith(key[:4]) and key[2:] in t.replace(" ", ""):
                # map to lead-in paragraph instead of heading
                lead = {
                    "3.3实体抽取与规范化": "在实体抽取与规范化方面，",
                    "3.4关系推理与边权更新": "在关系推理与边权更新方面，",
                    "3.5新兴岗位发现与预测": "在新兴岗位发现与预测方面，",
                    "3.6既有岗位能力动态演化": "在既有岗位能力动态演化方面，",
                    "3.7人岗匹配与学习路径": "在人岗匹配与学习路径方面，",
                    "3.8大模型对话与幻觉防控": "在大模型对话与幻觉防控方面，",
                    "3.9实验与结果": "实验环境与结果如下。",
                    "3.10总结与展望": "",
                }.get(key, "")
                # 下一正文段若存在，把 lead 并入下一段；本标题段删除
                nxt = p._element.getnext()
                if key.endswith("总结与展望"):
                    delete_paragraph(p)
                elif lead and nxt is not None:
                    nxt_p = Paragraph(nxt, p._parent)
                    body = nxt_p.text.strip()
                    body = strip_label_colon(body)
                    if not body.startswith(lead[:4]):
                        set_paragraph_text(nxt_p, lead + body)
                    delete_paragraph(p)
                else:
                    delete_paragraph(p)
                break

    # 图谱形式化：去掉可能的标签；该段在 3.3 总起之后
    p_graph = find_para(startswith="定义能力知识图谱")
    if p_graph is not None:
        set_paragraph_text(p_graph, "图谱形式化方面，" + strip_label_colon(p_graph.text))

    # 去掉消融实验两段
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("消融实验") or ("消融降幅依次" in t and "技能（0.42）" in t):
            delete_paragraph(p)

    # 清理“问题形式化与六阶段推理链：”“三维新兴度评分：”等标签体
    label_prefixes = (
        "问题形式化与六阶段推理链：",
        "问题形式化与六阶段推理链:",
        "三维新兴度评分：",
        "三维新兴度评分:",
        "岗位定义生成与未来方向修正：",
        "岗位定义生成与未来方向修正:",
        "证据审计：",
        "证据审计:",
        "技能抽取与出现率：",
        "技能抽取与出现率:",
        "输出与人工复核：",
        "输出与人工复核:",
        "问题形式化定义：",
        "问题形式化定义:",
        "五维加权模型与符号：",
        "五维加权模型与符号:",
        "差距分析与学习路径：",
        "差距分析与学习路径:",
        "生成约束与降级：",
        "生成约束与降级:",
        "实验环境与数据集：",
        "实验环境与数据集:",
        "创新点—证据对照：",
        "创新点—证据对照:",
        "工作总结：",
        "工作总结:",
        "局限：",
        "局限:",
    )
    for p in doc.paragraphs:
        t = p.text.strip()
        for lab in label_prefixes:
            if t.startswith(lab):
                rest = t[len(lab):].strip()
                # 对工作总结去掉消融表述
                if "消融验证" in rest:
                    rest = rest.replace("，消融验证权重合理性", "")
                set_paragraph_text(p, rest)
                break

    # 工作总结、局限再润色（若仍带标签已被剥掉）
    for p in doc.paragraphs:
        t = p.text
        if "本文面向赛题XH-202621" in t and "消融" in t:
            set_paragraph_text(p, t.replace("，消融验证权重合理性", ""))
        if t.startswith("配对t检验") or ("去掉技能覆盖" in t and "配对t检验" in t):
            # leftover ablation residue
            delete_paragraph(p)

    # 统一所有 Heading 2 字体为黑体
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 2":
            text = p.text.strip()
            set_heading_text(p, text, 2)
        if p.style and p.style.name == "Heading 1":
            text = p.text.strip()
            set_heading_text(p, text, 1)

    # 图题保持黑体居中感（不强制居中以免破坏版式）
    for p in doc.paragraphs:
        t = p.text.strip()
        if re.match(r"^图\d+", t) or re.match(r"^表\d+", t):
            clear_paragraph(p)
            run = p.add_run(t)
            set_run_east_asia(run, "黑体", 10.5, False)
            p.paragraph_format.first_line_indent = None
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(DST))
    remove_all_comments(DST)
    print("saved", DST)


if __name__ == "__main__":
    main()
