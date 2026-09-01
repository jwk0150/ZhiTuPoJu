# -*- coding: utf-8 -*-
"""将 作品设计实现方案.txt 导出为 Word (.docx)。"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
TXT = ROOT / "作品设计实现方案.txt"
DOCX = ROOT / "作品设计实现方案.docx"

CHAPTER = re.compile(r"^第[一二三四五六七八九十]+章\s")
SEC1 = re.compile(r"^\d+\.\d+\s")
SEC2 = re.compile(r"^\d+\.\d+\.\d+\s")
SEC_I = re.compile(r"^2\.\d+\s")
TABLE_ROW = re.compile(r"^\t.+\t")


def set_cn_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, align=None, bold=False, size=12):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    return p


def parse_table_block(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith("【表"):
            continue
        parts = [c.strip() for c in line.split("\t")]
        if len(parts) >= 2:
            rows.append(parts)
    return rows


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < cols:
                table.rows[i].cells[j].text = cell
    doc.add_paragraph()


def main():
    text = TXT.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    i = 0
    table_buf: list[str] = []
    in_table = False
    title_done = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("【表开始】"):
            in_table = True
            table_buf = []
            i += 1
            continue
        if stripped.startswith("【表结束】"):
            in_table = False
            add_table(doc, parse_table_block(table_buf))
            table_buf = []
            i += 1
            continue
        if in_table:
            table_buf.append(line)
            i += 1
            continue

        if not stripped or stripped.startswith("="):
            i += 1
            continue

        if not title_done and ("作品设计实现方案" in stripped or "执图破局" in stripped):
            if "作品设计" in stripped or stripped.startswith("执图破局"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                set_cn_font(run, name="黑体", size=22, bold=True)
                title_done = True
                i += 1
                continue

        if stripped in ("摘要", "关键词", "目录"):
            add_para(doc, stripped, bold=True, size=14)
            i += 1
            continue

        if stripped.startswith("关键词：") or stripped.startswith("副标题："):
            add_para(doc, stripped, size=12)
            i += 1
            continue

        if CHAPTER.match(stripped):
            doc.add_heading(stripped, level=1)
            i += 1
            continue
        if SEC2.match(stripped):
            doc.add_heading(stripped, level=3)
            i += 1
            continue
        if SEC1.match(stripped) or SEC_I.match(stripped):
            doc.add_heading(stripped, level=2)
            i += 1
            continue

        if stripped.startswith("参考文献") or stripped.startswith("附录"):
            doc.add_heading(stripped, level=1)
            i += 1
            continue

        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt or nxt.startswith("=") or nxt.startswith("【表"):
                break
            if CHAPTER.match(nxt) or SEC1.match(nxt) or SEC2.match(nxt) or SEC_I.match(nxt):
                break
            if nxt in ("摘要", "关键词", "目录") or nxt.startswith("参考文献") or nxt.startswith("附录"):
                break
            para_lines.append(nxt)
            j += 1
        add_para(doc, "\n".join(para_lines), size=12)
        i = j

    doc.save(DOCX)
    print(f"OK: {DOCX}")


if __name__ == "__main__":
    main()
