# -*- coding: utf-8 -*-
"""
直接生成《作品设计实现方案》Word（国赛计划书体例封面 + 论文正文）。
参照：石墨烯基柔性锂离子电容器计划书封面结构（口号页眉、英文导语、大标题、组别信息）。

运行：python scripts/devtools/generate_design_doc.py
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from design_doc.ch01_overview import SECTIONS as CH01
from design_doc.ch02_core import SECTIONS as CH02
from design_doc.ch03_architecture import SECTIONS as CH03
from design_doc.ch04_innovation import SECTIONS as CH04
from design_doc.ch05_testing import SECTIONS as CH05
from design_doc.ch06_backmatter import ABSTRACT, APPENDICES, REFERENCES, SECTIONS as CH06
from design_doc.formatting import (
    add_body,
    add_cover_image_section,
    add_figure_caption,
    add_heading,
    add_page_break,
    add_table,
    add_table_caption,
    add_toc_field,
    configure_document,
    set_run_font,
)

OUT = ROOT / "作品设计实现方案_提交版.docx"
COVER_SCRIPT = Path(__file__).resolve().parent / "design_doc" / "compose_cover.py"
COVER_PNG = Path(__file__).resolve().parent / "design_doc" / "assets" / "cover-composed.png"
ATMOS_PNG = Path(__file__).resolve().parent / "design_doc" / "assets" / "cover-atmosphere.png"


def ensure_cover():
    subprocess.check_call([sys.executable, str(COVER_SCRIPT)])
    if not COVER_PNG.exists():
        raise FileNotFoundError(COVER_PNG)


def render_sections(doc, sections):
    chars = 0
    for sec in sections:
        lvl = sec.get("level", 2)
        title = sec.get("title", "")
        if title:
            add_heading(doc, title, level=lvl)
        if sec.get("figure"):
            add_figure_caption(doc, sec["figure"])
        tbl = sec.get("table")
        if tbl:
            if tbl.get("caption"):
                add_table_caption(doc, tbl["caption"])
            add_table(doc, tbl["headers"], tbl["rows"])
        for para in sec.get("paragraphs", []):
            add_body(doc, para)
            chars += len(para)
    return chars


def exec_summary_page(doc):
    """封面后执行概要页（参照国赛计划书开篇结构）。"""
    add_heading(doc, "执行概要", 1)
    lead = (
        "执图破局面向赛题 XH-202621，针对岗位—能力关系分散、新兴岗位难识别、人岗匹配难解释等痛点，"
        "构建「多源异构数据驱动的岗位—能力知识图谱动态构建与智能匹配系统」。平台以可信语料入库、"
        "新兴岗位发现、能力演化差分、五维匹配与 RAG 幻觉防控为闭环，形成可演示、可评测、可复现的完整作品。"
    )
    add_body(doc, lead)
    bullets = [
        "目标用户：高校学生、求职者、就业研究者与平台运营方；以「发现—演化—匹配」三条主路径交付价值。",
        "核心能力：I1 入库门控、I2 新兴岗位发现、I3 演化差分、I4 五维人岗匹配、I5 RAG 幻觉防控。",
        "技术路线：PostgreSQL 多库视图 + FastAPI 服务编排 + DeepSeek 推理增强 + 前端一屏座舱可视化。",
        "交付形态：可运行前后端系统、作品设计实现方案、演示脚本与评测证据链。",
    ]
    for b in bullets:
        add_body(doc, b)
    if ATMOS_PNG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(ATMOS_PNG), width=Cm(14.5))
        add_figure_caption(doc, "图 0.1  执图破局：岗位—能力图谱与人岗连接意象")
    add_page_break(doc)


def abstract_page(doc):
    add_heading(doc, "摘要", 1)
    chars = 0
    for para in ABSTRACT["paragraphs"]:
        add_body(doc, para)
        chars += len(para)
    p = doc.add_paragraph()
    set_run_font(p.add_run("关键词："), bold=True)
    set_run_font(p.add_run(ABSTRACT["keywords"]))
    add_page_break(doc)
    return chars


def chapter2_intro(doc):
    add_heading(doc, "第二章  核心技术", 1)
    add_body(
        doc,
        "执图破局平台的核心竞争力源于对前沿数据治理、智能体推理与多维匹配技术的深度整合与创新应用。"
        "平台以「多源可信语料」为基石，致力于打破传统关键词匹配与静态图谱的局限性，通过整合入库门控、"
        "新兴岗位发现、能力演化差分、五维人岗匹配与 RAG 幻觉防控等多维度能力，为赛题 XH-202621 "
        "提供科学、客观、全面且可复现的智能化解决方案。",
    )
    add_figure_caption(doc, "图 2.1  执图破局核心技术体系")
    add_body(
        doc,
        "如图 2.1 所示，系统通过发现—演化—匹配协同分析，突破单一数据源局限，实现更全面的岗位—能力图谱构建。"
        "尤其在新兴岗位识别可回放性、演化差分可复算性以及匹配结果可解释性上表现突出，形成五项可定位创新（I1—I5）。",
    )


def references_page(doc):
    add_page_break(doc)
    add_heading(doc, "参考文献", 1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        set_run_font(p.add_run(ref), size=10.5)
        p.paragraph_format.first_line_indent = None


def count_all_sections():
    all_secs = CH01 + CH02 + CH03 + CH04 + CH05 + CH06 + APPENDICES
    return sum(len("".join(s.get("paragraphs", []))) for s in all_secs) + sum(
        len(p) for p in ABSTRACT["paragraphs"]
    )


def main():
    total_chars = count_all_sections()
    print(f"内容字数（段落合计）：约 {total_chars} 字")

    print("合成封面…")
    ensure_cover()

    doc = Document()
    configure_document(doc)
    add_cover_image_section(doc, COVER_PNG)
    exec_summary_page(doc)
    abstract_page(doc)

    add_heading(doc, "目录", 1)
    add_toc_field(doc)
    add_page_break(doc)

    render_sections(doc, CH01)
    add_page_break(doc)

    chapter2_intro(doc)
    render_sections(doc, [s for s in CH02 if s.get("level", 1) > 1])
    add_page_break(doc)

    render_sections(doc, CH03)
    add_page_break(doc)

    render_sections(doc, CH04)
    add_page_break(doc)

    render_sections(doc, CH05)
    add_page_break(doc)

    render_sections(doc, CH06)
    references_page(doc)
    render_sections(doc, APPENDICES)

    doc.save(OUT)
    print(f"已生成：{OUT}")
    print("请在 Word 中右键目录 → 更新域，以刷新页码。")


if __name__ == "__main__":
    main()
