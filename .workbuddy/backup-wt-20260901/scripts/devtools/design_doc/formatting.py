# -*- coding: utf-8 -*-
"""Word 论文体例排版（参照 word-thesis-format + 国赛计划书体例）。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HEADER_LEFT = "「执图破局」助力揭榜"
HEADER_RIGHT = "让人才与机会更精准连接"
FOOTER_BAR = "执图破局 平台设计方案"
ASSETS = Path(__file__).resolve().parent / "assets"
CORNER_PNG = ASSETS / "page-corner.png"
HEADER_BANNER_PNG = ASSETS / "page-header-banner.png"
FOOTER_BAR_PNG = ASSETS / "page-footer-bar.png"
PAGE_BG_PNG = ASSETS / "page-bg.png"
# 亮青绿底条（避免近黑压抑 + 文字被裁切）
FOOTER_FILL = "2A8F9A"
TEAL = (42, 143, 154)


def configure_document(doc: Document):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
    for lvl, size in [(1, 16), (2, 14), (3, 12)]:
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "黑体"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hr = hs.element.get_or_add_rPr()
        hf = hr.find(qn("w:rFonts"))
        if hf is None:
            hf = OxmlElement("w:rFonts")
            hr.insert(0, hf)
        hf.set(qn("w:eastAsia"), "黑体")


def set_run_font(run, name="宋体", size=12, bold=False, en="Times New Roman", color=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), en)


def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in kwargs.items():
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _emu(cm_val: float) -> int:
    """厘米 → EMU（1 inch = 914400 EMU）。"""
    return int(cm_val * 914400 / 2.54)


def add_page_float_picture(
    paragraph,
    image_path: Path,
    *,
    width_cm: float,
    height_cm: float | None = None,
    left_cm: float,
    top_cm: float,
    behind: bool = True,
    name: str = "PageOrnament",
):
    """
    页锚定浮动图：相对整页定位，不被页眉高度裁切。
    先插 inline，再改写成 wp:anchor。
    """
    from docx.oxml import parse_xml

    run = paragraph.add_run()
    if height_cm is None:
        inline = run.add_picture(str(image_path), width=Cm(width_cm))
    else:
        inline = run.add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))

    inline_el = inline._inline
    extent = inline_el.extent
    cx, cy = extent.cx, extent.cy
    graphic = inline_el.xpath("./a:graphic")[0]

    behind_flag = "1" if behind else "0"
    anchor_xml = (
        f'<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        f'distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" '
        f'behindDoc="{behind_flag}" locked="0" layoutInCell="1" allowOverlap="1">'
        f"<wp:simplePos x=\"0\" y=\"0\"/>"
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{_emu(left_cm)}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{_emu(top_cm)}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f"<wp:wrapNone/>"
        f'<wp:docPr id="{abs(hash(name)) % 100000}" name="{name}"/>'
        f"<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect=\"1\"/></wp:cNvGraphicFramePr>"
        f"</wp:anchor>"
    )
    anchor = parse_xml(anchor_xml)
    anchor.append(graphic)

    drawing = inline_el.getparent()
    drawing.remove(inline_el)
    drawing.append(anchor)
    return anchor


def add_header_footer(section, include_header=True):
    """善泽风格页装：每页浅色背景 + 顶栏横幅 + 亮青绿项目底条。"""
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.4)

    if include_header:
        header = section.header
        header.is_linked_to_previous = False
        for p in list(header.paragraphs):
            p.clear()
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)

        # 每页满版浅底（页锚定、衬于文字后）
        if PAGE_BG_PNG.exists():
            add_page_float_picture(
                hp,
                PAGE_BG_PNG,
                width_cm=21.0,
                height_cm=29.7,
                left_cm=0.0,
                top_cm=0.0,
                behind=True,
                name="ZhituPageBg",
            )

        banner = HEADER_BANNER_PNG if HEADER_BANNER_PNG.exists() else CORNER_PNG
        if banner.exists():
            run = hp.add_run()
            run.add_picture(str(banner), width=Cm(16.8))

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p.clear()
    fp0 = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp0.paragraph_format.space_before = Pt(0)
    fp0.paragraph_format.space_after = Pt(0)
    fp0.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 单行底条：左文案（加宽防截断）+ 右页码
    table = footer.add_table(rows=1, cols=2, width=Cm(16.8))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    c0, c1 = table.rows[0].cells
    # 明确列宽，避免「XH-202」被裁成半截
    c0.width = Cm(13.6)
    c1.width = Cm(3.2)
    _shade_cell(c0, FOOTER_FILL)
    _shade_cell(c1, FOOTER_FILL)
    _set_cell_margins(c0, top=85, bottom=85, left=160, right=60)
    _set_cell_margins(c1, top=85, bottom=85, left=40, right=120)

    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run(FOOTER_BAR)
    set_run_font(r0, name="微软雅黑", size=11, bold=True, en="微软雅黑", color=(255, 255, 255))

    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p1.add_run()
    set_run_font(run, name="微软雅黑", size=11, bold=True, en="微软雅黑", color=(255, 255, 255))
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_body_section(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    # 顶栏约 4cm，顶距留足，避免裁切
    section.top_margin = Cm(4.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    add_header_footer(section, include_header=True)


def add_body(doc, text, indent=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(
            run,
            name="黑体",
            size={1: 16, 2: 14, 3: 12}.get(level, 12),
            bold=True,
            en="微软雅黑",
            color=TEAL if level == 1 else (26, 92, 102),
        )
    h.paragraph_format.first_line_indent = None
    h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), size=10.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    set_run_font(p.add_run(text), size=10.5, bold=True)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        _shade_cell(cell, "E8F5F7")
        for para in cell.paragraphs:
            for r in para.runs:
                set_run_font(r, name="黑体", size=10.5, bold=True, en="微软雅黑", color=TEAL)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            if i % 2 == 1:
                _shade_cell(cell, "F7FBFC")
            for para in cell.paragraphs:
                for r in para.runs:
                    set_run_font(r, size=10.5)
    doc.add_paragraph()


def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)


def add_page_break(doc):
    doc.add_page_break()


def add_cover_image_section(doc, cover_png):
    """封面：满版 A4 合成图 + 新节进入正文。"""
    sec0 = doc.sections[0]
    sec0.page_height = Cm(29.7)
    sec0.page_width = Cm(21.0)
    sec0.top_margin = Cm(0)
    sec0.bottom_margin = Cm(0)
    sec0.left_margin = Cm(0)
    sec0.right_margin = Cm(0)
    for part in (sec0.header, sec0.footer):
        part.is_linked_to_previous = False
        if part.paragraphs:
            part.paragraphs[0].clear()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(cover_png), width=Cm(21.0), height=Cm(29.7))

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_body_section(body)
    return body
